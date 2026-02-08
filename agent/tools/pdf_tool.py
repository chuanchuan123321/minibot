"""PDF generation tool for converting various formats to PDF"""
import os
from pathlib import Path
from typing import Tuple, Dict, Any
from agent.tools.file import FileTool


class PDFTool:
    """Tool for generating PDF from various formats"""

    @staticmethod
    def _get_chinese_font():
        """Get Chinese font path for the system"""
        import platform
        system = platform.system()

        # Try to find Chinese fonts
        font_paths = []

        if system == "Darwin":  # macOS
            font_paths = [
                "/Library/Fonts/SimHei.ttf",
                "/Library/Fonts/SimSun.ttf",
                "/System/Library/Fonts/PingFang.ttc",
                "/Library/Fonts/Arial Unicode.ttf",
            ]
        elif system == "Linux":
            font_paths = [
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            ]
        elif system == "Windows":
            font_paths = [
                "C:\\Windows\\Fonts\\simhei.ttf",
                "C:\\Windows\\Fonts\\simsun.ttc",
                "C:\\Windows\\Fonts\\arial.ttf",
            ]

        # Find first available font
        for font_path in font_paths:
            if os.path.exists(font_path):
                return font_path

        return None

    @staticmethod
    def _register_chinese_font():
        """Register Chinese font with reportlab"""
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            font_path = PDFTool._get_chinese_font()
            if font_path:
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    return 'ChineseFont'
                except Exception as e:
                    print(f"Warning: Failed to register font {font_path}: {e}")

            return None
        except Exception as e:
            print(f"Warning: Failed to register Chinese font: {e}")
            return None

    @staticmethod
    def markdown_to_pdf(md_path: str, output_path: str) -> Tuple[bool, str]:
        """Convert Markdown file to PDF with proper formatting using fpdf2"""
        try:
            import markdown
            from html.parser import HTMLParser
        except ImportError as e:
            if "markdown" in str(e):
                return False, "Error: markdown not installed. Try: pip install markdown"
            else:
                return False, f"Error: Missing dependency - {str(e)}"

        try:
            # Try to use fpdf2 for better rendering
            try:
                from fpdf import FPDF
                use_fpdf = True
            except ImportError:
                use_fpdf = False

            expanded_path = FileTool.expand_path(md_path)
            if not os.path.isfile(expanded_path):
                return False, f"Error: File not found: {md_path}"

            with open(expanded_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # Parse markdown to HTML
            html = markdown.markdown(md_content, extensions=['extra', 'codehilite'])

            expanded_output = FileTool.expand_path(output_path)
            os.makedirs(os.path.dirname(expanded_output), exist_ok=True)

            if use_fpdf:
                # Use fpdf2 for better rendering
                class MarkdownPDF(FPDF):
                    def __init__(self):
                        super().__init__()
                        self.set_auto_page_break(auto=True, margin=15)
                        # Try to add Chinese font
                        self.font_name = 'Arial'
                        try:
                            font_path = PDFTool._get_chinese_font()
                            if font_path:
                                self.add_font('Chinese', '', font_path)
                                self.font_name = 'Chinese'
                        except:
                            pass

                    def add_title(self, text):
                        try:
                            self.set_font(self.font_name, 'B', 18)
                        except:
                            self.set_font(self.font_name, '', 18)
                        self.set_text_color(31, 71, 136)
                        self.cell(0, 10, text, ln=True)
                        self.ln(5)

                    def add_heading2(self, text):
                        try:
                            self.set_font(self.font_name, 'B', 14)
                        except:
                            self.set_font(self.font_name, '', 14)
                        self.set_text_color(46, 92, 138)
                        self.cell(0, 8, text, ln=True)
                        self.ln(3)

                    def add_heading3(self, text):
                        try:
                            self.set_font(self.font_name, 'B', 12)
                        except:
                            self.set_font(self.font_name, '', 12)
                        self.set_text_color(61, 111, 163)
                        self.cell(0, 7, text, ln=True)
                        self.ln(2)

                    def add_paragraph(self, text):
                        self.set_font(self.font_name, '', 11)
                        self.set_text_color(0, 0, 0)
                        self.multi_cell(0, 5, text)
                        self.ln(3)

                    def add_bullet(self, text):
                        self.set_font(self.font_name, '', 11)
                        self.set_text_color(0, 0, 0)
                        self.cell(5, 5, '•')
                        # Use set_x to position the text after the bullet
                        x_pos = self.get_x() + 5
                        self.set_x(x_pos)
                        self.multi_cell(0, 5, text)
                        self.ln(1)

                    def add_code(self, text):
                        self.set_font('Courier', '', 9)
                        self.set_fill_color(245, 245, 245)
                        self.set_text_color(0, 0, 0)
                        for line in text.split('\n'):
                            self.cell(0, 5, line, ln=True, fill=True)
                        self.ln(2)

                # Parse HTML and extract elements
                class MarkdownHTMLParser(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.elements = []
                        self.current_text = []
                        self.current_tag = None

                    def handle_starttag(self, tag, attrs):
                        if tag in ['h1', 'h2', 'h3', 'p', 'li', 'pre']:
                            self._flush_text()
                            self.current_tag = tag

                    def handle_endtag(self, tag):
                        if tag == self.current_tag:
                            text = ''.join(self.current_text).strip()
                            if text:
                                self.elements.append((tag, text))
                            self.current_text = []
                            self.current_tag = None

                    def handle_data(self, data):
                        self.current_text.append(data)

                    def _flush_text(self):
                        text = ''.join(self.current_text).strip()
                        if text and self.current_tag:
                            self.elements.append((self.current_tag, text))
                        self.current_text = []

                    def get_elements(self):
                        self._flush_text()
                        return self.elements

                parser = MarkdownHTMLParser()
                parser.feed(html)
                elements = parser.get_elements()

                # Create PDF and add elements
                pdf = MarkdownPDF()
                pdf.add_page()

                for tag, text in elements:
                    if tag == 'h1':
                        pdf.add_title(text)
                    elif tag == 'h2':
                        pdf.add_heading2(text)
                    elif tag == 'h3':
                        pdf.add_heading3(text)
                    elif tag == 'p':
                        pdf.add_paragraph(text)
                    elif tag == 'li':
                        pdf.add_bullet(text)
                    elif tag == 'pre':
                        pdf.add_code(text)

                pdf.output(expanded_output)

            else:
                # Fallback to reportlab
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import inch
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

                doc = SimpleDocTemplate(expanded_output, pagesize=A4)
                story = []

                font_name = PDFTool._register_chinese_font()
                styles = getSampleStyleSheet()

                if font_name:
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontName=font_name,
                        fontSize=18,
                        leading=22,
                        spaceAfter=12,
                        textColor='#1f4788',
                    )
                    heading2_style = ParagraphStyle(
                        'CustomHeading2',
                        parent=styles['Heading2'],
                        fontName=font_name,
                        fontSize=14,
                        leading=18,
                        spaceAfter=10,
                        textColor='#2e5c8a',
                    )
                    normal_style = ParagraphStyle(
                        'CustomNormal',
                        parent=styles['Normal'],
                        fontName=font_name,
                        fontSize=11,
                        leading=14,
                    )
                else:
                    title_style = styles['Heading1']
                    heading2_style = styles['Heading2']
                    normal_style = styles['Normal']

                class MarkdownHTMLParser(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.elements = []
                        self.current_text = []
                        self.current_tag = None

                    def handle_starttag(self, tag, attrs):
                        if tag in ['h1', 'h2', 'h3', 'p', 'li', 'pre']:
                            self._flush_text()
                            self.current_tag = tag

                    def handle_endtag(self, tag):
                        if tag == self.current_tag:
                            text = ''.join(self.current_text).strip()
                            if text:
                                self.elements.append((tag, text))
                            self.current_text = []
                            self.current_tag = None

                    def handle_data(self, data):
                        self.current_text.append(data)

                    def _flush_text(self):
                        text = ''.join(self.current_text).strip()
                        if text and self.current_tag:
                            self.elements.append((self.current_tag, text))
                        self.current_text = []

                    def get_elements(self):
                        self._flush_text()
                        return self.elements

                parser = MarkdownHTMLParser()
                parser.feed(html)
                elements = parser.get_elements()

                for tag, text in elements:
                    if tag == 'h1':
                        story.append(Paragraph(text, title_style))
                        story.append(Spacer(1, 0.2 * inch))
                    elif tag == 'h2':
                        story.append(Paragraph(text, heading2_style))
                        story.append(Spacer(1, 0.15 * inch))
                    elif tag == 'p':
                        story.append(Paragraph(text, normal_style))
                        story.append(Spacer(1, 0.1 * inch))
                    elif tag == 'li':
                        story.append(Paragraph(f"• {text}", normal_style))
                    elif tag == 'pre':
                        for line in text.split('\n'):
                            story.append(Paragraph(line, normal_style))
                        story.append(Spacer(1, 0.1 * inch))

                doc.build(story)

            file_size = os.path.getsize(expanded_output)
            return True, f"Markdown converted to PDF: {expanded_output} ({file_size} bytes)"

        except Exception as e:
            return False, f"Error converting Markdown to PDF: {str(e)}"

    @staticmethod
    def text_to_pdf(text_path: str, output_path: str) -> Tuple[bool, str]:
        """Convert text file to PDF using fpdf2"""
        try:
            from fpdf import FPDF
        except ImportError:
            return False, "Error: fpdf2 not installed. Try: pip install fpdf2"

        try:
            expanded_path = FileTool.expand_path(text_path)
            if not os.path.isfile(expanded_path):
                return False, f"Error: File not found: {text_path}"

            # Detect encoding
            try:
                import chardet
                with open(expanded_path, 'rb') as f:
                    raw_data = f.read()
                    result = chardet.detect(raw_data)
                    encoding = result.get('encoding', 'utf-8') or 'utf-8'
            except ImportError:
                encoding = 'utf-8'

            with open(expanded_path, 'r', encoding=encoding) as f:
                text_content = f.read()

            expanded_output = FileTool.expand_path(output_path)
            os.makedirs(os.path.dirname(expanded_output), exist_ok=True)

            # Create PDF using fpdf2
            class TextPDF(FPDF):
                def __init__(self):
                    super().__init__()
                    self.set_auto_page_break(auto=True, margin=15)
                    self.font_name = 'Arial'
                    try:
                        font_path = PDFTool._get_chinese_font()
                        if font_path:
                            self.add_font('Chinese', '', font_path)
                            self.font_name = 'Chinese'
                    except:
                        pass

            pdf = TextPDF()
            pdf.add_page()
            pdf.set_font(pdf.font_name, '', 11)
            pdf.set_text_color(0, 0, 0)

            # Add text content
            for line in text_content.split('\n'):
                if line.strip():
                    pdf.multi_cell(0, 5, line)
                else:
                    pdf.ln(3)

            pdf.output(expanded_output)

            file_size = os.path.getsize(expanded_output)
            return True, f"Text converted to PDF: {expanded_output} ({file_size} bytes)"

        except Exception as e:
            return False, f"Error converting text to PDF: {str(e)}"

    @staticmethod
    def html_to_pdf(html_path: str, output_path: str) -> Tuple[bool, str]:
        """Convert HTML file to PDF using fpdf2"""
        try:
            from fpdf import FPDF
        except ImportError:
            return False, "Error: fpdf2 not installed. Try: pip install fpdf2"

        try:
            expanded_path = FileTool.expand_path(html_path)
            if not os.path.isfile(expanded_path):
                return False, f"Error: File not found: {html_path}"

            with open(expanded_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Use BeautifulSoup if available, otherwise use HTMLParser
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')

                # Remove script and style tags
                for script in soup(["script", "style"]):
                    script.decompose()

                # Get text
                text = soup.get_text()

                # Clean up text
                lines = []
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        lines.append(line)
                text = '\n'.join(lines)

            except ImportError:
                # Fallback to HTMLParser
                from html.parser import HTMLParser

                class HTMLTextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text = []
                        self.skip_tags = {'script', 'style', 'meta', 'link', 'head'}
                        self.skip_level = 0

                    def handle_starttag(self, tag, attrs):
                        if tag in self.skip_tags:
                            self.skip_level += 1

                    def handle_endtag(self, tag):
                        if tag in self.skip_tags:
                            self.skip_level = max(0, self.skip_level - 1)

                    def handle_data(self, data):
                        if self.skip_level == 0:
                            text = data.strip()
                            if text:
                                self.text.append(text + ' ')

                    def get_text(self):
                        return ''.join(self.text).strip()

                extractor = HTMLTextExtractor()
                extractor.feed(html_content)
                text = extractor.get_text()

            expanded_output = FileTool.expand_path(output_path)
            os.makedirs(os.path.dirname(expanded_output), exist_ok=True)

            # Create PDF using fpdf2
            class HTMLPDF(FPDF):
                def __init__(self):
                    super().__init__()
                    self.set_auto_page_break(auto=True, margin=15)
                    self.font_name = 'Arial'
                    try:
                        font_path = PDFTool._get_chinese_font()
                        if font_path:
                            self.add_font('Chinese', '', font_path)
                            self.font_name = 'Chinese'
                    except:
                        pass

            pdf = HTMLPDF()
            pdf.add_page()
            pdf.set_font(pdf.font_name, '', 11)
            pdf.set_text_color(0, 0, 0)

            # Add text content
            for line in text.split('\n'):
                if line.strip():
                    pdf.multi_cell(0, 5, line)
                else:
                    pdf.ln(2)

            pdf.output(expanded_output)

            file_size = os.path.getsize(expanded_output)
            return True, f"HTML converted to PDF: {expanded_output} ({file_size} bytes)"

        except Exception as e:
            return False, f"Error converting HTML to PDF: {str(e)}"

    @staticmethod
    def docx_to_pdf(docx_path: str, output_path: str) -> Tuple[bool, str]:
        """Convert DOCX file to PDF using fpdf2 with style support"""
        try:
            from docx import Document
            from fpdf import FPDF
        except ImportError as e:
            if "docx" in str(e):
                return False, "Error: python-docx not installed. Try: pip install python-docx"
            elif "fpdf" in str(e):
                return False, "Error: fpdf2 not installed. Try: pip install fpdf2"
            else:
                return False, f"Error: Missing dependency - {str(e)}"

        try:
            expanded_path = FileTool.expand_path(docx_path)
            if not os.path.isfile(expanded_path):
                return False, f"Error: File not found: {docx_path}"

            doc = Document(expanded_path)

            expanded_output = FileTool.expand_path(output_path)
            os.makedirs(os.path.dirname(expanded_output), exist_ok=True)

            # Create PDF using fpdf2
            class DocxPDF(FPDF):
                def __init__(self):
                    super().__init__()
                    self.set_auto_page_break(auto=True, margin=10)
                    self.font_name = 'Arial'
                    try:
                        font_path = PDFTool._get_chinese_font()
                        if font_path:
                            self.add_font('Chinese', '', font_path)
                            self.font_name = 'Chinese'
                    except:
                        pass

                def add_heading1(self, text):
                    """Add Heading 1 style"""
                    self.set_font(self.font_name, '', 16)
                    self.set_text_color(31, 71, 136)
                    self.write(6, text + "\n")
                    self.set_text_color(0, 0, 0)
                    self.ln(2)

                def add_heading2(self, text):
                    """Add Heading 2 style"""
                    self.set_font(self.font_name, '', 13)
                    self.set_text_color(46, 92, 138)
                    self.write(5, text + "\n")
                    self.set_text_color(0, 0, 0)
                    self.ln(1)

                def add_heading3(self, text):
                    """Add Heading 3 style"""
                    self.set_font(self.font_name, '', 11)
                    self.set_text_color(61, 111, 163)
                    self.write(4, text + "\n")
                    self.set_text_color(0, 0, 0)
                    self.ln(1)

                def add_normal(self, text):
                    """Add normal paragraph"""
                    self.set_font(self.font_name, '', 10)
                    self.set_text_color(0, 0, 0)
                    self.write(4, text + "\n")
                    self.ln(1)

                def add_list_item(self, text):
                    """Add list item with bullet"""
                    self.set_font(self.font_name, '', 10)
                    self.set_text_color(0, 0, 0)
                    self.write(4, "• " + text + "\n")
                    self.ln(0.5)

            pdf = DocxPDF()
            pdf.add_page()

            # Process paragraphs with style detection
            for para in doc.paragraphs:
                if not para.text.strip():
                    pdf.ln(1)
                    continue

                text = para.text.replace('\t', '    ')
                style_name = para.style.name if para.style else ""

                # Detect heading levels
                if "Heading 1" in style_name or style_name.startswith("标题 1"):
                    pdf.add_heading1(text)
                elif "Heading 2" in style_name or style_name.startswith("标题 2"):
                    pdf.add_heading2(text)
                elif "Heading 3" in style_name or style_name.startswith("标题 3"):
                    pdf.add_heading3(text)
                elif "List" in style_name or para.paragraph_format.left_indent:
                    # Detect list items by style or indentation
                    pdf.add_list_item(text)
                else:
                    pdf.add_normal(text)

            # Process tables
            for table in doc.tables:
                pdf.ln(2)
                pdf.set_font(pdf.font_name, '', 9)
                pdf.write(3, f"[表格: {len(table.rows)} 行 × {len(table.columns)} 列]\n")
                pdf.ln(1)

                for row in table.rows:
                    pdf.set_font(pdf.font_name, '', 9)
                    row_text = " | ".join([cell.text.replace('\t', ' ') for cell in row.cells])
                    pdf.write(3, row_text + "\n")
                pdf.ln(1)

            pdf.output(expanded_output)

            file_size = os.path.getsize(expanded_output)
            return True, f"DOCX converted to PDF: {expanded_output} ({file_size} bytes)"

        except Exception as e:
            return False, f"Error converting DOCX to PDF: {str(e)}"

    @staticmethod
    def generate_pdf(input_path: str, output_path: str, format_type: str) -> Tuple[bool, str]:
        """Generate PDF from various formats"""
        if format_type == "markdown":
            return PDFTool.markdown_to_pdf(input_path, output_path)
        elif format_type == "text":
            return PDFTool.text_to_pdf(input_path, output_path)
        elif format_type == "html":
            return PDFTool.html_to_pdf(input_path, output_path)
        elif format_type == "docx":
            return PDFTool.docx_to_pdf(input_path, output_path)
        else:
            return False, f"Error: Unsupported format '{format_type}'. Supported: markdown, text, html, docx"

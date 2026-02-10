"""Extended tool executor with document reading capabilities"""
import json
from typing import Dict, Any, Callable, Optional
from agent.tools.shell import ShellTool
from agent.tools.file import FileTool
from agent.tools.time_tool import TimeTool
from agent.tools.pdf_tool import PDFTool
from agent.tools.skill_tool import SkillTool
import os
import requests


class ExtendedToolExecutor:
    """Execute tools with extended capabilities including document reading"""

    def __init__(self, skills_loader=None):
        self.shell_tool = ShellTool()
        self.file_tool = FileTool()
        self.pdf_tool = PDFTool()
        self.skill_tool = SkillTool(skills_loader) if skills_loader else None
        self.tools: Dict[str, Callable] = {
            "shell": self.execute_shell,
            "file_read": self.execute_file_read,
            "file_write": self.execute_file_write,
            "file_list": self.execute_file_list,
            "file_delete": self.execute_file_delete,
            "dir_create": self.execute_dir_create,
            "dir_change": self.execute_dir_change,
            "read_pdf": self.execute_read_pdf,
            "read_markdown": self.execute_read_markdown,
            "read_json": self.execute_read_json,
            "search_files": self.execute_search_files,
            "get_file_info": self.execute_get_file_info,
            "copy_file": self.execute_copy_file,
            "move_file": self.execute_move_file,
            "create_file": self.execute_create_file,
            "web_search": self.execute_web_search,
            "read_url": self.execute_read_url,
            "set_timer": self.execute_set_timer,
            "send_file": self.execute_send_file,
            "generate_pdf": self.execute_generate_pdf,
            "load_skill": self.execute_load_skill,
        }

    def get_available_tools(self) -> list:
        """Get list of available tools"""
        return [
            {
                "name": "shell",
                "description": "Execute shell commands on the system",
                "params": "command (string): The shell command to execute"
            },
            {
                "name": "file_read",
                "description": "Read the contents of a text file",
                "params": "path (string): Path to the file to read"
            },
            {
                "name": "file_write",
                "description": "Write content to a file",
                "params": "path (string): File path, content (string): Content to write"
            },
            {
                "name": "file_list",
                "description": "List files in a directory",
                "params": "path (string): Directory path (default: current directory)"
            },
            {
                "name": "file_delete",
                "description": "Delete a file",
                "params": "path (string): Path to the file to delete"
            },
            {
                "name": "dir_create",
                "description": "Create a directory",
                "params": "path (string): Path to the directory to create"
            },
            {
                "name": "dir_change",
                "description": "Change the current working directory",
                "params": "path (string): Path to change to"
            },
            {
                "name": "read_pdf",
                "description": "Read and extract text from PDF files",
                "params": "path (string): Path to the PDF file"
            },
            {
                "name": "read_markdown",
                "description": "Read and parse markdown files",
                "params": "path (string): Path to the markdown file"
            },
            {
                "name": "read_json",
                "description": "Read and parse JSON files",
                "params": "path (string): Path to the JSON file"
            },
            {
                "name": "search_files",
                "description": "Search for files by name or pattern",
                "params": "pattern (string): File name pattern to search for, path (string): Directory to search in"
            },
            {
                "name": "get_file_info",
                "description": "Get detailed information about a file",
                "params": "path (string): Path to the file"
            },
            {
                "name": "copy_file",
                "description": "Copy a file to a new location",
                "params": "source (string): Source file path, destination (string): Destination path"
            },
            {
                "name": "move_file",
                "description": "Move or rename a file",
                "params": "source (string): Source file path, destination (string): Destination path"
            },
            {
                "name": "create_file",
                "description": "Create a new file with content",
                "params": "path (string): File path, content (string): File content"
            },
            {
                "name": "web_search",
                "description": "Search the web for information",
                "params": "query (string): Search query"
            },
            {
                "name": "read_url",
                "description": "Read and extract content from a URL",
                "params": "url (string): The URL to read"
            },
            {
                "name": "set_timer",
                "description": "Set a timer that will trigger after specified minutes",
                "params": "minutes (number): Minutes to wait, message (string): Message to display when timer ends"
            },
            {
                "name": "send_file",
                "description": "Send a file to the user via Feishu",
                "params": "path (string): Path to the file to send"
            },
            {
                "name": "generate_pdf",
                "description": "Generate PDF from Markdown, text, HTML, or Word documents",
                "params": "input_path (string): Input file path, output_path (string): Output PDF file path, format (string): Input format (markdown/text/html/docx)"
            },
            {
                "name": "load_skill",
                "description": "Load a skill's complete content to get detailed guidance and instructions",
                "params": "skill_name (string): Name of the skill to load (e.g., 'web', 'github', 'python')"
            },
        ]

    def execute(self, tool_call: Dict[str, Any]) -> str:
        """Execute a tool call"""
        tool_name = tool_call.get("tool")
        params = tool_call.get("params", {})

        if tool_name not in self.tools:
            return f"Error: Unknown tool '{tool_name}'"

        try:
            result = self.tools[tool_name](params)
            return result
        except Exception as e:
            return f"Error executing {tool_name}: {str(e)}"

    def execute_shell(self, params: Dict[str, Any]) -> str:
        """Execute shell command"""
        command = params.get("command", "")
        if not command:
            return "Error: command parameter required"

        result = self.shell_tool.execute(command)
        return self.shell_tool.format_result(result)

    def execute_file_read(self, params: Dict[str, Any]) -> str:
        """Read file"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        success, content = self.file_tool.read_file(path)
        if success:
            return f"File contents:\n{content}"
        return f"Error: {content}"

    def execute_file_write(self, params: Dict[str, Any]) -> str:
        """Write file"""
        path = params.get("path", "")
        content = params.get("content", "")
        if not path:
            return "Error: path parameter required"

        success, message = self.file_tool.write_file(path, content)
        return message if success else f"Error: {message}"

    def execute_file_list(self, params: Dict[str, Any]) -> str:
        """List files"""
        path = params.get("path", ".")
        success, files = self.file_tool.list_files(path)

        if success:
            file_list = "\n".join(files[:50])
            if len(files) > 50:
                file_list += f"\n... and {len(files) - 50} more files"
            return f"Files in {path}:\n{file_list}"
        return f"Error: {files[0] if files else 'Unknown error'}"

    def execute_file_delete(self, params: Dict[str, Any]) -> str:
        """Delete file"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        success, message = self.file_tool.delete_file(path)
        return message if success else f"Error: {message}"

    def execute_dir_create(self, params: Dict[str, Any]) -> str:
        """Create directory"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        success, message = self.file_tool.create_directory(path)
        return message if success else f"Error: {message}"

    def execute_dir_change(self, params: Dict[str, Any]) -> str:
        """Change directory"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        success, message = self.shell_tool.change_dir(path)
        return message if success else f"Error: {message}"

    def execute_read_pdf(self, params: Dict[str, Any]) -> str:
        """Read PDF or document file"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        try:
            expanded_path = FileTool.expand_path(path)

            # 检查文件类型
            if expanded_path.endswith('.docx') or expanded_path.endswith('.doc'):
                # 处理Word文档
                try:
                    from docx import Document
                    doc = Document(expanded_path)
                    text = ""
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"

                    # 也提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text for cell in row.cells])
                            text += row_text + "\n"

                    return f"Document contents:\n{text}"
                except ImportError:
                    return "Error: python-docx not installed. Try: pip install python-docx"

            elif expanded_path.endswith('.pdf'):
                # 处理PDF文件
                try:
                    import PyPDF2
                    with open(expanded_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in reader.pages:  # Read all pages
                            text += page.extract_text() + "\n"
                    return f"PDF contents:\n{text}"
                except ImportError:
                    return "Error: PyPDF2 not installed. Try: pip install PyPDF2"

            else:
                return f"Error: Unsupported file format. Supported: .pdf, .docx, .doc"

        except Exception as e:
            return f"Error reading document: {str(e)}"

    def execute_read_markdown(self, params: Dict[str, Any]) -> str:
        """Read markdown file"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        success, content = self.file_tool.read_file(path)
        if success:
            return f"Markdown contents:\n{content[:2000]}"
        return f"Error: {content}"

    def execute_read_json(self, params: Dict[str, Any]) -> str:
        """Read and parse JSON file"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        try:
            success, content = self.file_tool.read_file(path)
            if success:
                data = json.loads(content)
                return f"JSON contents:\n{json.dumps(data, indent=2, ensure_ascii=False)[:2000]}"
            return f"Error: {content}"
        except json.JSONDecodeError as e:
            return f"Error: Invalid JSON format - {str(e)}"

    def execute_search_files(self, params: Dict[str, Any]) -> str:
        """Search for files"""
        pattern = params.get("pattern", "")
        path = params.get("path", ".")
        if not pattern:
            return "Error: pattern parameter required"

        try:
            result = self.shell_tool.execute(f"find {FileTool.expand_path(path)} -name '*{pattern}*' -type f | head -20")
            if result.success:
                return f"Found files:\n{result.stdout}"
            return f"No files found matching pattern: {pattern}"
        except Exception as e:
            return f"Error searching files: {str(e)}"

    def execute_get_file_info(self, params: Dict[str, Any]) -> str:
        """Get file information"""
        path = params.get("path", "")
        if not path:
            return "Error: path parameter required"

        success, info = self.file_tool.get_file_info(path)
        if success:
            return f"File info:\n{json.dumps(info, indent=2, ensure_ascii=False)}"
        return f"Error: {info.get('error', 'Unknown error')}"

    def execute_copy_file(self, params: Dict[str, Any]) -> str:
        """Copy file"""
        source = params.get("source", "")
        destination = params.get("destination", "")
        if not source or not destination:
            return "Error: source and destination parameters required"

        try:
            import shutil
            source_path = FileTool.expand_path(source)
            dest_path = FileTool.expand_path(destination)
            shutil.copy2(source_path, dest_path)
            return f"File copied: {source} -> {destination}"
        except Exception as e:
            return f"Error copying file: {str(e)}"

    def execute_move_file(self, params: Dict[str, Any]) -> str:
        """Move or rename file"""
        source = params.get("source", "")
        destination = params.get("destination", "")
        if not source or not destination:
            return "Error: source and destination parameters required"

        try:
            import shutil
            source_path = FileTool.expand_path(source)
            dest_path = FileTool.expand_path(destination)
            shutil.move(source_path, dest_path)
            return f"File moved: {source} -> {destination}"
        except Exception as e:
            return f"Error moving file: {str(e)}"

    def execute_create_file(self, params: Dict[str, Any]) -> str:
        """Create file with content"""
        path = params.get("path", "")
        content = params.get("content", "")
        if not path:
            return "Error: path parameter required"

        success, message = self.file_tool.write_file(path, content)
        return message if success else f"Error: {message}"

    def execute_web_search(self, params: Dict[str, Any]) -> str:
        """Search the web using Tavily API"""
        query = params.get("query", "")
        if not query:
            return "Error: query parameter required"

        try:
            tavily_api_key = os.getenv("TAVILY_API_KEY")
            if not tavily_api_key:
                return "Error: TAVILY_API_KEY not found in environment variables"

            search_url = "https://api.tavily.com/search"
            payload = {
                "api_key": tavily_api_key,
                "query": query,
                "include_answer": True,
                "max_results": 5
            }

            response = requests.post(search_url, json=payload, timeout=10)
            response.raise_for_status()
            data = response.json()

            results = []

            # Add answer if available
            if data.get("answer"):
                results.append(f"答案: {data['answer']}")
                results.append("")

            # Add search results
            if data.get("results"):
                results.append("搜索结果:")
                for result in data.get("results", [])[:10]:
                    if result.get("title"):
                        results.append(f"- {result['title']}")
                    if result.get("content"):
                        results.append(f"  {result['content'][:150]}...")
                    if result.get("url"):
                        results.append(f"  链接: {result['url']}")
                    results.append("")

            if results:
                return "搜索结果:\n" + "\n".join(results)
            else:
                return f"未找到关于 '{query}' 的搜索结果"

        except requests.exceptions.Timeout:
            return "Error: 搜索请求超时"
        except requests.exceptions.RequestException as e:
            return f"Error: 网络请求失败 - {str(e)}"
        except Exception as e:
            return f"Error: 搜索失败 - {str(e)}"

    def execute_read_url(self, params: Dict[str, Any]) -> str:
        """Read and extract content from a URL"""
        url = params.get("url", "")
        if not url:
            return "Error: url parameter required"

        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            }

            response = requests.get(url, headers=headers, timeout=15)

            # Try to detect and set correct encoding
            if response.encoding is None or response.encoding.lower() == 'iso-8859-1':
                # Try to detect encoding from content
                try:
                    import chardet
                    detected = chardet.detect(response.content)
                    if detected and detected.get('encoding'):
                        response.encoding = detected['encoding']
                    else:
                        response.encoding = 'utf-8'
                except ImportError:
                    # If chardet not available, try common encodings
                    for encoding in ['utf-8', 'gb2312', 'gbk', 'big5', 'iso-8859-1']:
                        try:
                            response.content.decode(encoding)
                            response.encoding = encoding
                            break
                        except:
                            continue

            response.raise_for_status()

            # Get text content
            content = response.text

            # Try using BeautifulSoup if available
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(content, 'html.parser')

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Get text
                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)

                if text:
                    return f"URL 内容:\n{text[:5000]}"
                else:
                    return "Error: 无法从 URL 提取内容"

            except ImportError:
                # If BeautifulSoup not available, try simple regex extraction
                import re

                # Remove script and style tags
                content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
                content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)

                # Remove HTML tags
                content = re.sub(r'<[^>]+>', '', content)

                # Clean up whitespace
                content = re.sub(r'\s+', ' ', content).strip()

                if content:
                    return f"URL 内容:\n{content[:5000]}"
                else:
                    return "Error: 无法从 URL 提取内容"

        except requests.exceptions.Timeout:
            return "Error: 请求超时"
        except requests.exceptions.RequestException as e:
            return f"Error: 网络请求失败 - {str(e)}"
        except Exception as e:
            return f"Error: 读取 URL 失败 - {str(e)}"

    def execute_set_timer(self, params: Dict[str, Any]) -> str:
        """Set a timer that will trigger after specified minutes"""
        import time
        import threading

        minutes = params.get("minutes", 0)
        message = params.get("message", "时间到了！")
        executor = params.get("executor", None)  # 获取执行器引用

        if not isinstance(minutes, (int, float)) or minutes <= 0:
            return "Error: minutes 必须是正数"

        try:
            seconds = minutes * 60

            def timer_callback():
                """Timer callback function"""
                time.sleep(seconds)
                print(f"\n⏰ 【定时器触发】{message}\n")

                # 如果有执行器引用，设置标志
                if executor:
                    executor.timer_triggered = True
                    executor.waiting_for_timer = False

            # 在后台线程中运行定时器
            timer_thread = threading.Thread(target=timer_callback, daemon=True)
            timer_thread.start()

            return f"✅ 定时器已设置：{minutes}分钟后将显示 '{message}'"

        except Exception as e:
            return f"Error: 设置定时器失败 - {str(e)}"

    def execute_send_file(self, params: Dict[str, Any]) -> str:
        """Send a file to the user via Feishu"""
        import asyncio
        from agent.bus.events import OutboundMessage

        file_path = params.get("path", "")

        if not file_path:
            return "Error: 必须指定文件路径"

        import os
        if not os.path.isfile(file_path):
            return f"Error: 文件不存在 - {file_path}"

        # Get the current executor context to access bus and chat info
        # This is a bit hacky but necessary for the current architecture
        import inspect
        frame = inspect.currentframe()
        executor = None

        # Walk up the stack to find NaturalTaskExecutor
        while frame:
            if 'self' in frame.f_locals:
                obj = frame.f_locals['self']
                if hasattr(obj, 'bus') and hasattr(obj, 'current_chat_id'):
                    executor = obj
                    break
            frame = frame.f_back

        if not executor or not executor.bus or not executor.current_chat_id:
            return "Error: 无法发送文件 - 未在网关模式下运行"

        try:
            # Create outbound message with file path
            msg = OutboundMessage(
                channel=executor.current_channel or "feishu",
                chat_id=executor.current_chat_id,
                content=file_path,  # Pass file path as content
            )

            # Send via bus (non-blocking)
            asyncio.create_task(executor.bus.publish_outbound(msg))

            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            return f"✅ 文件已发送: {file_name} ({file_size} bytes)"

        except Exception as e:
            return f"Error: 发送文件失败 - {str(e)}"

    def execute_generate_pdf(self, params: Dict[str, Any]) -> str:
        """Generate PDF from various formats"""
        input_path = params.get("input_path", "")
        output_path = params.get("output_path", "")
        format_type = params.get("format", "")

        # Parameter validation
        if not input_path or not output_path:
            return "Error: input_path and output_path parameters required"

        # Auto-detect format from file extension if not specified
        if not format_type:
            if input_path.lower().endswith('.md'):
                format_type = "markdown"
            elif input_path.lower().endswith('.html'):
                format_type = "html"
            elif input_path.lower().endswith('.docx') or input_path.lower().endswith('.doc'):
                format_type = "docx"
            else:
                format_type = "text"

        if format_type not in ["markdown", "text", "html", "docx"]:
            return f"Error: Unsupported format '{format_type}'. Supported: markdown, text, html, docx"

        # Path expansion
        expanded_input = FileTool.expand_path(input_path)
        expanded_output = FileTool.expand_path(output_path)

        # Call PDF tool
        success, message = self.pdf_tool.generate_pdf(
            expanded_input,
            expanded_output,
            format_type
        )

        if success:
            return f"✅ {message}"
        else:
            return message

    def execute_load_skill(self, params: Dict[str, Any]) -> str:
        """Load a skill's complete content"""
        if not self.skill_tool:
            return "Error: Skill tool not initialized"

        skill_name = params.get("skill_name", "").strip()
        if not skill_name:
            return "Error: skill_name parameter required"

        success, content = self.skill_tool.load_skill(skill_name)
        return content

